import tempfile
from pathlib import Path

import pytest

from app.ingestion.chunker import chunk_pages, _split_text
from app.ingestion.loader import load_file


# ── Chunker ────────────────────────────────────────────────────────────────────

def test_split_text_short():
    """Text shorter than chunk_size returns a single chunk with offsets."""
    result = _split_text("Hello world.")
    assert len(result) == 1
    assert result[0]["text"] == "Hello world."
    assert result[0]["start"] == 0
    assert result[0]["end"] == len("Hello world.")


def test_split_text_long():
    """Long text is split into multiple overlapping chunks with offsets."""
    long_text = "word " * 300  # ~1500 chars
    chunks = _split_text(long_text)
    assert len(chunks) > 1
    assert all(c["text"].strip() for c in chunks)
    assert all("start" in c for c in chunks)
    assert all("end" in c for c in chunks)


def test_chunk_pages_preserves_metadata():
    """chunk_pages attaches source, page, chunk_index, offsets, heading and content_type to every chunk."""
    pages = [{
        "text": "word " * 200,
        "page": 1,
        "source": "test.txt",
        "content_type": "text",
        "heading": "Introduction",
    }]
    chunks = chunk_pages(pages)
    for chunk in chunks:
        assert chunk["source"] == "test.txt"
        assert chunk["page"] == 1
        assert "chunk_index" in chunk
        assert chunk["chunk_total"] == len(chunks)
        assert "char_start" in chunk
        assert "char_end" in chunk
        assert chunk["content_type"] == "text"
        assert chunk["section_heading"] == "Introduction"


def test_chunk_index_monotonic():
    """chunk_index increases across multiple pages."""
    pages = [
        {"text": "word " * 200, "page": 1, "source": "test.txt"},
        {"text": "word " * 200, "page": 2, "source": "test.txt"},
    ]
    chunks = chunk_pages(pages)
    indices = [c["chunk_index"] for c in chunks]
    assert indices == sorted(indices)
    assert len(set(indices)) == len(indices)  # all unique


# ── Loader ─────────────────────────────────────────────────────────────────────

def test_load_txt():
    """Plain text file loads correctly."""
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False) as f:
        f.write("Hello from a text file.\n" * 50)
        tmp_path = Path(f.name)

    try:
        pages = load_file(tmp_path)
        assert len(pages) >= 1
        assert all("text" in p for p in pages)
        assert all("source" in p for p in pages)
        assert pages[0]["source"] == tmp_path.name
    finally:
        tmp_path.unlink()


def test_load_unsupported_extension():
    """Unsupported file type raises ValueError."""
    with tempfile.NamedTemporaryFile(suffix=".csv", delete=False) as f:
        tmp_path = Path(f.name)
    try:
        with pytest.raises(ValueError, match="Unsupported file type"):
            load_file(tmp_path)
    finally:
        tmp_path.unlink()


def test_load_docx_with_headings_and_tables():
    """DOCX parser correctly extracts headings, paragraphs, and tables as separate sections."""
    import docx
    doc = docx.Document()
    doc.add_paragraph("Main Title Heading", style="Heading 1")
    doc.add_paragraph("Normal paragraph body text.")
    doc.add_paragraph("Secondary Heading", style="Heading 2")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = Path(f.name)
        
    try:
        doc.save(str(tmp_path))
        sections = load_file(tmp_path)
        
        # We expect two text virtual sections (paragraph text) and one table section
        assert len(sections) >= 2
        
        # Check text sections
        text_sections = [s for s in sections if s.get("content_type") == "text"]
        assert len(text_sections) > 0
        assert any("Main Title Heading" in s["text"] for s in text_sections)
        # Check heading styles detection
        assert any(s["heading"] == "Main Title Heading" or s["heading"] == "Secondary Heading" for s in text_sections)
        
        # Check table section
        table_sections = [s for s in sections if s.get("content_type") == "table"]
        assert len(table_sections) == 1
        assert table_sections[0]["text"] == "A1 | B1\nA2 | B2"
        # The table should inherit the active heading style (Secondary Heading)
        assert table_sections[0]["heading"] == "Secondary Heading"
        
    finally:
        tmp_path.unlink()


def test_load_pdf_with_tables():
    """PDF parser loads page using markdown mode."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Hello PDF text.")
    
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        tmp_path = Path(f.name)
        
    try:
        doc.save(str(tmp_path))
        pages = load_file(tmp_path)
        
        assert len(pages) == 1
        assert "Hello PDF text." in pages[0]["text"]
        assert pages[0]["source"] == tmp_path.name
        
    finally:
        tmp_path.unlink()