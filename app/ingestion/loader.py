from pathlib import Path
import fitz  # PyMuPDF
import docx

from config import settings


def load_file(file_path: Path) -> list[dict]:
    """
    Load a file and return a list of page/section dicts:
        [{"text": str, "page": int, "source": str}, ...]

    Supports .pdf, .docx, .txt
    """
    suffix = file_path.suffix.lower()

    if suffix not in settings.allowed_extensions:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Allowed: {settings.allowed_extensions}"
        )

    if suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix == ".docx":
        return _load_docx(file_path)
    elif suffix == ".txt":
        return _load_txt(file_path)


# ── Private loaders ────────────────────────────────────────────────────────────

def _load_pdf(file_path: Path) -> list[dict]:
    """Extract text page-by-page from a PDF using PyMuPDF."""
    pages = []
    with fitz.open(str(file_path)) as doc:
        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks")
            # Sort blocks to read top-to-bottom, left-to-right
            blocks.sort(key=lambda b: (b[1], b[0]))
            text = "\n\n".join(b[4].strip() for b in blocks if b[4].strip())
            if text:  # skip blank pages
                pages.append({
                    "text": text,
                    "page": page_num,
                    "source": file_path.name,
                })
    return pages


def _load_docx(file_path: Path) -> list[dict]:
    """
    Extract text and tables from a DOCX file.
    Groups paragraphs into virtual 'pages' of ~2000 chars to keep
    metadata granular without producing 1-sentence chunks.
    Also extracts tables and detects headings.
    """
    doc = docx.Document(str(file_path))

    sections = []
    current_text = ""
    virtual_page = 1
    current_heading = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style and para.style.name and para.style.name.startswith("Heading"):
            current_heading = text

        current_text += text + "\n"
        if len(current_text) >= 2000:
            sections.append({
                "text": current_text.strip(),
                "page": virtual_page,
                "source": file_path.name,
                "content_type": "text",
                "heading": current_heading,
            })
            current_text = ""
            virtual_page += 1

    # Flush remainder
    if current_text.strip():
        sections.append({
            "text": current_text.strip(),
            "page": virtual_page,
            "source": file_path.name,
            "content_type": "text",
            "heading": current_heading,
        })
        virtual_page += 1

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            sections.append({
                "text": table_text.strip(),
                "page": virtual_page,
                "source": file_path.name,
                "content_type": "table",
                "heading": current_heading,
            })
            virtual_page += 1

    return sections


def _load_txt(file_path: Path) -> list[dict]:
    """
    Load a plain text file.
    Splits into virtual pages every 3000 characters.
    """
    text = file_path.read_text(encoding="utf-8", errors="replace")
    page_size = 3000
    pages = []

    for i, start in enumerate(range(0, len(text), page_size), start=1):
        chunk_text = text[start : start + page_size].strip()
        if chunk_text:
            pages.append({
                "text": chunk_text,
                "page": i,
                "source": file_path.name,
            })

    return pages